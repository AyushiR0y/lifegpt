"""
docx_generator.py
Creates a professional Word (.docx) document from a plain-text summary.

Formatting rules
----------------
• Section headers (standalone short lines / ## / ###) → BOLD, coloured
• "Label: Description" lines → Label BOLD, description normal
• Bullet lines (•, -, *) with "Label: Description" → same bold split
• Regular text → normal weight
• No ** or __ markdown is expected in the input (already cleaned upstream)
"""

import io
from datetime import datetime


def generate_simple_summary_docx(
    summary_text: str,
    original_filename: str,
    summary_type: str,
) -> bytes:
    """
    Convert a plain-text *summary_text* to a formatted Word document.

    Parameters
    ----------
    summary_text      : Summary produced by the LLM (no ** markdown).
    original_filename : Name of the source document (shown in metadata line).
    summary_type      : One of ``concise``, ``mid_level``, ``descriptive``.

    Returns
    -------
    bytes : Raw DOCX file content ready for streaming / download.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ------------------------------------------------------------------
    # Base style – everything starts as plain non-bold Calibri 11
    # ------------------------------------------------------------------
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.bold = False
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)

    # ------------------------------------------------------------------
    # Title block
    # ------------------------------------------------------------------
    _add_run(
        doc.add_paragraph(),
        "Document Summary",
        bold=True, size=22,
        color=RGBColor(0, 51, 102),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4),
    )

    meta_text = (
        f"Source: {original_filename}  |  "
        f"Type: {summary_type.replace('_', ' ').title()}  |  "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    _add_run(
        doc.add_paragraph(),
        meta_text,
        bold=False, size=9,
        color=RGBColor(128, 128, 128),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(2),
    )

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after = Pt(6)
    r = sep.add_run("─" * 60)
    r.font.color.rgb = RGBColor(180, 180, 180)

    # ------------------------------------------------------------------
    # Process each line of the summary
    # ------------------------------------------------------------------
    for raw_line in summary_text.split("\n"):
        line = raw_line.strip()

        # Strip residual markdown artefacts just in case
        for tok in ("**", "__", "```"):
            line = line.replace(tok, "")

        # --- Empty line → small spacer paragraph ---
        if not line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            continue

        # --- Markdown H1 (# …) ---
        if line.startswith("# ") and not line.startswith("## "):
            _header_para(doc, line[2:].strip(), size=14, color=RGBColor(0, 51, 102),
                         space_before=10, space_after=3)
            continue

        # --- Markdown H2 (## …) ---
        if line.startswith("## "):
            _header_para(doc, line[3:].strip(), size=13, color=RGBColor(0, 76, 153),
                         space_before=8, space_after=2)
            continue

        # --- Markdown H3 (### …) ---
        if line.startswith("### "):
            _header_para(doc, line[4:].strip(), size=11, color=RGBColor(51, 51, 51),
                         space_before=6, space_after=2)
            continue

        # --- Table / box-drawing characters ---
        if line.startswith("|") or (line and line[0] in "┌├└│"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if line.startswith("|"):
                p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.bold = False
            continue

        # --- Separator lines ---
        if line.startswith("─") or line.startswith("---") or line.startswith("==="):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run("─" * 40)
            r.font.color.rgb = RGBColor(180, 180, 180)
            r.bold = False
            continue

        # --- Standalone section header (short, no colon, starts with capital) ---
        if _looks_like_header(line):
            _header_para(doc, line, size=12, color=RGBColor(0, 76, 153),
                         space_before=8, space_after=2)
            continue

        # --- Bullet point ---
        if line.startswith(("•", "·", "- ", "* ")):
            bullet_text = line.lstrip("•·-* ").strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _label_desc(p, bullet_text)
            continue

        # --- Non-bullet "Label: Description" ---
        colon = line.find(":")
        if 0 < colon < 50 and not line.lower().startswith("http"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _label_desc_full(p, line[:colon].strip(), line[colon + 1:].strip())
            continue

        # --- Regular text ---
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        r.bold = False
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # ------------------------------------------------------------------
    # Footer
    # ------------------------------------------------------------------
    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(10)
    fr = fp.add_run("─" * 60)
    fr.font.color.rgb = RGBColor(180, 180, 180)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Generated by AI Document Summarizer")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(128, 128, 128)
    r.bold = False

    # ------------------------------------------------------------------
    # Serialise
    # ------------------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _looks_like_header(line: str) -> bool:
    """True when the line looks like a standalone section title."""
    if ":" in line:
        return False
    if len(line) >= 50 or len(line.split()) > 6:
        return False
    if line.startswith(("•", "·", "-", "|", "─")) or (line and line[0].isdigit()):
        return False
    words = line.split()
    return bool(words) and words[0][0].isupper()


def _header_para(
    doc,
    text: str,
    *,
    size: int,
    color,
    space_before: int,
    space_after: int,
) -> None:
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color


def _label_desc(para, bullet_text: str) -> None:
    """Split 'Label: description' inside a bullet paragraph."""
    from docx.shared import Pt
    colon = bullet_text.find(":")
    if 0 < colon < 50:
        label = bullet_text[:colon].strip()
        desc = bullet_text[colon + 1:].strip()
        r = para.add_run(label + ": ")
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        if desc:
            r2 = para.add_run(desc)
            r2.bold = False
            r2.font.size = Pt(11)
            r2.font.name = "Calibri"
    else:
        r = para.add_run(bullet_text)
        r.bold = False
        r.font.size = Pt(11)
        r.font.name = "Calibri"


def _label_desc_full(para, label: str, desc: str) -> None:
    """Bold label + normal description for non-bullet lines."""
    from docx.shared import Pt
    r = para.add_run(label + ": ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    if desc:
        r2 = para.add_run(desc)
        r2.bold = False
        r2.font.size = Pt(11)
        r2.font.name = "Calibri"


def _add_run(
    para,
    text: str,
    *,
    bold: bool,
    size: int,
    color,
    align=None,
    space_after=None,
) -> None:
    from docx.shared import Pt
    if align is not None:
        para.alignment = align
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
